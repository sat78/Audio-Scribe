
import streamlit as st
import base64
from pathlib import Path
import io
import logging
from dotenv import load_dotenv
import os
import tempfile
from datetime import datetime
from moviepy import VideoFileClip
from pydub import AudioSegment
from pydub.utils import make_chunks
import speech_recognition as sr
import json
from chatbot import get_insights_from_video
from fastapi import FastAPI
from pydantic import BaseModel
from typing import Optional


try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import assemblyai as aai
    ASSEMBLYAI_AVAILABLE = True
except ImportError:
    ASSEMBLYAI_AVAILABLE = False

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False


# load_dotenv('./.env')
# logger.info(f"GOOGLE_API_KEY: {os.getenv('GOOGLE_API_KEY')}")
# Page configuration (MUST be first Streamlit command)
st.set_page_config(
    page_title="AudioScribe - AI Transcription",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.session_state["MAX_UPLOAD_SIZE_MB"] = 500
st.markdown("<style>.stFileUploader input[type=file] {max-file-size: 500MB;}</style>", unsafe_allow_html=True)


def get_base64_image(image_path):
    """Convert image to base64"""
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except FileNotFoundError:
        st.warning(f"⚠️ Background image not found at: {image_path}")
        return None
    except Exception as e:
        st.warning(f"⚠️ Error loading background: {e}")
        return None

import os

import requests
BACKEND_URL = "http://127.0.0.1:8009/chat"

def ask_backend(question, transcription):
    payload = {
        "query": question,
        "transcription": transcription
    }
    res = requests.post(BACKEND_URL, json=payload)
    return res.json().get("answer", "No answer")


# Get the directory where the script is located
script_dir = os.path.dirname(os.path.abspath(__file__))

# OPTION 1: Use assets folder
background_path = os.path.join(script_dir, "assets", "podcast.jpg")

# OPTION 2: If assets folder doesn't work, try direct path
# background_path = "assets/background.png"

# Debug output
print(f"Script directory: {script_dir}")
print(f"Background path: {background_path}")
print(f"File exists: {os.path.exists(background_path)}")

background_b64 = get_base64_image(background_path)


if background_b64:
    st.markdown(f"""
        <style>
        #MainMenu {{visibility: hidden;}}
        footer {{visibility: hidden;}}
        header {{visibility: hidden;}}
        
        /* Background on root with darker overlay for better contrast */
        .stApp {{
            background: linear-gradient(rgba(255, 255, 255, 0.35), rgba(255, 255, 255, 0.35)),
                url('data:image/png;base64,{background_b64}') !important;
    background-size: cover !important;
    background-position: center !important;
    background-attachment: fixed !important;
    background-repeat: no-repeat !important;
}}

        
        /* Transparent main containers */
        .main {{
            background: transparent !important;
        }}
        
        .main .block-container {{
            background: transparent !important;
        }}
        
        /* Header - Keep solid */
        .main-header {{
            background: linear-gradient(90deg, #009688 0%, #263238 100%) !important;
            padding: 1rem 2rem;
            margin: -6rem -6rem 2rem -6rem;
            color: white;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
            position: relative;
            z-index: 100;
        }}
        
        .header-content {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            max-width: 1400px;
            margin: 0 auto;
        }}
        
        .logo-section {{
            display: flex;
            align-items: center;
            gap: 1rem;
        }}
        
        .logo-circle {{
            width: 40px;
            height: 40px;
            background: white;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: bold;
            color: #009688;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.2);
        }}
        
        /* ALL HEADINGS - WHITE with shadow */
        .main .block-container h1,
        .main .block-container h2,
        .main .block-container h3,
        .main .block-container h4 {{
            color: #ffffff !important;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.8);
            font-weight: 600 !important;
        }}
        
        /* ALL PARAGRAPHS - black */
        .main .block-container p {{
            color: #000000 !important;
            text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.7);
        }}
        
        /* ALL DIV TEXT - WHITE */
        .main .block-container div {{
            color: #000000 !important;
        }}
        
        /* MARKDOWN TEXT - WHITE */
        .main .block-container .markdown-text-container {{
            color: #000000 !important;
        }}
        
        /* SIDEBAR HEADINGS - WHITE with teal accent */
        [data-testid="stSidebar"] h1,
        [data-testid="stSidebar"] h2,
        [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] h4 {{
            color: #000000 !important;
            text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.6);
        }}
        
        /* SIDEBAR TEXT - WHITE */
        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] div,
        [data-testid="stSidebar"] span {{
            color: #000000 !important;
            text-shadow: 1px 1px 2px rgba(0, 0, 0, 0.5);
        }}
        
        /* SIDEBAR BACKGROUND - Semi-transparent */
        [data-testid="stSidebar"] {{
            background: rgba(10, 50, 50, 0.7) !important;
            backdrop-filter: blur(15px);
            box-shadow: 2px 0 12px rgba(0, 0, 0, 0.3);
        }}
        
        /* SIDEBAR section labels */
        [data-testid="stSidebar"] .element-container {{
            color: #000000 !important;
        }}
        
        /* Usage stats card - Keep white background, dark text */
        .usage-stats {{
            background: rgba(255, 255, 255, 0.98) !important;
            backdrop-filter: blur(12px);
            padding: 1rem;
            border-radius: 10px;
            margin-bottom: 1rem;
            border-left: 4px solid #009688;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.25);
        }}
        
        .usage-stats p {{
            color: #2d3748 !important;
            text-shadow: none !important;
        }}
        
        /* Cards - Solid white with dark text */
        .upload-modal {{
            background: rgba(255, 255, 255, 0.99) !important;
            backdrop-filter: blur(20px);
            border-radius: 12px;
            padding: 2rem;
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
            border: 1px solid rgba(0, 150, 136, 0.3);
        }}
        
        .upload-modal h1,
        .upload-modal h2,
        .upload-modal h3,
        .upload-modal p,
        .upload-modal div {{
            color: #2d3748 !important;
            text-shadow: none !important;
        }}
        
        /* Drag-drop area */
        .drag-drop-area {{
            border: 3px dashed #009688;
            border-radius: 12px;
            padding: 3rem;
            text-align: center;
            background: rgba(255, 255, 255, 0.98) !important;
            backdrop-filter: blur(10px);
            margin: 1.5rem 0;
            transition: all 0.3s ease;
            box-shadow: inset 0 2px 8px rgba(0, 0, 0, 0.05);
        }}
        
        .drag-drop-area:hover {{
            background: rgba(240, 255, 250, 0.99) !important;
            box-shadow: 0 4px 20px rgba(0, 150, 136, 0.3);
        }}
        
        /* Transcript container - Solid */
        .transcript-container {{
            background: rgba(255, 255, 255, 0.99) !important;
            backdrop-filter: blur(20px);
            border-radius: 12px;
            padding: 2rem;
            margin-top: 2rem;
            box-shadow: 0 4px 24px rgba(0, 0, 0, 0.25);
            border: 1px solid rgba(0, 150, 136, 0.2);
        }}
        
        .transcript-container * {{
            color: #2d3748 !important;
            text-shadow: none !important;
        }}
        
        /* File items */
        .file-item {{
            background: rgba(255, 255, 255, 0.98) !important;
            backdrop-filter: blur(10px);
            padding: 1rem;
            border-radius: 10px;
            margin-bottom: 0.5rem;
            border: 1px solid rgba(0, 150, 136, 0.25);
            box-shadow: 0 2px 12px rgba(0, 0, 0, 0.15);
            transition: all 0.3s ease;
        }}
        
        .file-item:hover {{
            background: rgba(255, 255, 255, 1) !important;
            box-shadow: 0 6px 24px rgba(0, 150, 136, 0.35);
            transform: translateX(4px);
        }}
        
        /* Buttons - More prominent */
        .stButton>button {{
            background: linear-gradient(135deg, #009688 0%, #00796B 100%) !important;
            color: #ffffff !important;
            border: none !important;
            border-radius: 10px !important;
            padding: 0.8rem 2rem !important;
            font-weight: 600 !important;
            transition: all 0.3s !important;
            box-shadow: 0 4px 16px rgba(0, 150, 136, 0.4) !important;
            text-shadow: 0 1px 2px rgba(0, 0, 0, 0.2);
        }}
        
        .stButton>button:hover {{
            background: linear-gradient(135deg, #00796B 0%, #00695C 100%) !important;
            box-shadow: 0 6px 20px rgba(0, 150, 136, 0.5) !important;
            transform: translateY(-2px);
        }}
        
        /* Timestamp styling */
        .timestamp {{
            color: #009688;
            font-weight: 700;
            font-size: 0.95em;
            margin-right: 0.5rem;
            background: rgba(0, 150, 136, 0.15);
            padding: 0.3rem 0.7rem;
            border-radius: 6px;
            border: 1px solid rgba(0, 150, 136, 0.3);
        }}
        
        /* Search results */
        .search-result-item {{
            background: rgba(255, 255, 255, 0.98) !important;
            backdrop-filter: blur(10px);
            padding: 0.8rem;
            margin: 0.5rem 0;
            border-radius: 8px;
            cursor: pointer;
            border: 1px solid rgba(0, 150, 136, 0.2);
            transition: all 0.3s ease;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        }}
        
        .search-result-item:hover {{
            background: rgba(240, 255, 250, 0.99) !important;
            border-color: #009688;
            box-shadow: 0 4px 16px rgba(0, 150, 136, 0.3);
        }}
        
        /* Input fields - Solid backgrounds */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea {{
            background: rgba(255, 255, 255, 0.98) !important;
            border: 1px solid rgba(0, 150, 136, 0.3) !important;
            color: #2d3748 !important;
        }}
        
        .stSelectbox > div > div {{
            background: rgba(255, 255, 255, 0.98) !important;
            border: 1px solid rgba(0, 150, 136, 0.3) !important;
        }}
        
        /* File uploader */
        [data-testid="stFileUploader"] {{
            background: rgba(255, 255, 255, 0.98) !important;
            border-radius: 10px;
            padding: 1rem;
        }}
        
        [data-testid="stFileUploader"] * {{
            color: #2d3748 !important;
        }}
        
        /* Progress bars */
        .stProgress > div > div > div {{
            background: linear-gradient(90deg, #009688 0%, #00796B 100%) !important;
        }}
        
        /* Info/Warning/Error boxes */
        .stAlert {{
            background: rgba(255, 255, 255, 0.98) !important;
            backdrop-filter: blur(10px);
            border-radius: 8px;
        }}
        
        /* Expander */
        .streamlit-expanderHeader {{
            background: rgba(255, 255, 255, 0.95) !important;
            border-radius: 8px;
        }}
        
        /* Caption text - keep visible */
        .caption {{
            color: #e0e0e0 !important;
            text-shadow: 1px 1px 2px rgba(0, 0, 0, 0.6) !important;
        }}
        
        /* Streamlit markdown - force black */
        .stMarkdown {{
            color: #000000 !important;
        }}
        </style>
    """, unsafe_allow_html=True)
# Initialize session state
if 'transcriptions' not in st.session_state:
    st.session_state.transcriptions = []
if 'current_transcription' not in st.session_state:
    st.session_state.current_transcription = None
if 'show_upload_modal' not in st.session_state:
    st.session_state.show_upload_modal = False
if 'show_sow_generator' not in st.session_state:
    st.session_state.show_sow_generator = False
if 'show_mom_generator' not in st.session_state:
    st.session_state.show_mom_generator = False
if 'generated_sow' not in st.session_state:
    st.session_state.generated_sow = None
if 'generated_mom' not in st.session_state:
    st.session_state.generated_mom = None
if 'show_search' not in st.session_state:
    st.session_state.show_search = False
if 'show_recent_files' not in st.session_state:
    st.session_state.show_recent_files = False
if 'gemini_api_key' not in st.session_state:
    st.session_state.gemini_api_key = ''
if 'assemblyai_key' not in st.session_state:
    # Replace this with your actual AssemblyAI API key
    st.session_state.assemblyai_key = 'b47ac0657d2a40a391296e6b09578629'

# def user_query(user_question, transcription_text):
#     """Wrapper to run async chatbot function in Streamlit"""
#     try:
#         loop = asyncio.new_event_loop()
#         asyncio.set_event_loop(loop)
#         result = loop.run_until_complete(
#             get_insights_from_video(user_question, transcription_text)
#         )
#         loop.close()
#         return result
#     except Exception as e:
#         logger.error(f"Error in user_query: {str(e)}")
#         return f"<h3 style='color:#ff0000;'>Error: {str(e)}</h3>"



# Enhanced transcription functions


def transcribe_video_to_text_enhanced(video_path):
    """Extract audio from video with optimized settings"""
    try:
        video_clip = VideoFileClip(video_path)
        audio_clip = video_clip.audio
        audio_path = "temp_audio.wav"
        
        audio_clip.write_audiofile(
            audio_path,
            fps=16000,
            nbytes=2,
            codec='pcm_s16le',
            logger=None,
            
        )
        
        audio_clip.close()
        video_clip.close()
        
        return audio_path
    except Exception as e:
        st.error(f"Error extracting audio: {e}")
        return None
    
    





def transcribe_with_assemblyai(audio_path, show_timestamps=False):
    """Super fast transcription using AssemblyAI"""
    if not ASSEMBLYAI_AVAILABLE:
        st.error("AssemblyAI not installed. Install it with: pip install assemblyai")
        return None, None
    
    try:
        status_text = st.empty()
        
        api_key = st.session_state.assemblyai_key
        
        if api_key == 'YOUR_ASSEMBLYAI_API_KEY_HERE':
            st.error("⚠️ Please set your AssemblyAI API key in the code (line 229)")
            st.info("Get free API key at: https://www.assemblyai.com/")
            return None, None
        
        status_text.text("☁️ Uploading to AssemblyAI...")
        
        aai.settings.api_key = api_key
        transcriber = aai.Transcriber()
        
        status_text.text("🎙️ Transcribing (this is fast!)...")
        
        config = aai.TranscriptionConfig(
            language_code="en",
            punctuate=True,
            format_text=True
        )
        
        transcript = transcriber.transcribe(audio_path, config=config)
        
        if transcript.status == aai.TranscriptStatus.error:
            st.error(f"Transcription failed: {transcript.error}")
            return None, None
        
        full_text = transcript.text
        timestamped_text = None
        
        if show_timestamps and hasattr(transcript, 'words'):
            timestamped_text = []
            current_segment = {"time": "", "text": ""}
            last_time = 0
            
            for word in transcript.words:
                time_seconds = word.start / 1000
                minutes = int(time_seconds // 60)
                seconds = int(time_seconds % 60)
                
                if time_seconds - last_time > 5:
                    if current_segment["text"]:
                        timestamped_text.append(current_segment.copy())
                    current_segment = {
                        "time": f"({minutes:02d}:{seconds:02d})",
                        "text": word.text
                    }
                    last_time = time_seconds
                else:
                    current_segment["text"] += " " + word.text
            
            if current_segment["text"]:
                timestamped_text.append(current_segment)
        
        status_text.text("✅ Transcription complete!")
        return full_text, timestamped_text
        
    except Exception as e:
        st.error(f"AssemblyAI error: {e}")
        return None, None



def transcribe_long_audio_enhanced(file_path, chunk_length_ms=60000, show_timestamps=False):
    """Enhanced Google API transcription with better error handling"""
    import time
    
    recognizer = sr.Recognizer()
    recognizer.energy_threshold = 300
    recognizer.dynamic_energy_threshold = True
    recognizer.pause_threshold = 0.8
    
    try:
        # Ensure proper WAV format
        audio = AudioSegment.from_file(file_path)
        audio = audio.set_channels(1).set_frame_rate(16000).set_sample_width(2)
        
        # Export to proper WAV
        temp_wav = "temp_google_api.wav"
        audio.export(temp_wav, format="wav")
        
        audio = AudioSegment.from_wav(temp_wav)
        chunks = make_chunks(audio, chunk_length_ms)
        
        full_text = ""
        timestamped_text = []
        failed_chunks = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks):
            chunk_filename = f"chunk_{i}.wav"
            
            # Export chunk with proper format
            chunk.set_channels(1).set_frame_rate(16000).export(chunk_filename, format="wav")
            
            status_text.text(f"🎙️ Processing chunk {i+1} of {len(chunks)}...")
            progress_bar.progress((i + 1) / len(chunks))
            
            with sr.AudioFile(chunk_filename) as source:
                audio_data = recognizer.record(source)
                
                max_retries = 3
                retry_delay = 2
                text = None
                
                for attempt in range(max_retries):
                    try:
                        text = recognizer.recognize_google(
                            audio_data, 
                            language="en-IN",
                            show_all=False
                        )
                        break
                        
                    except sr.UnknownValueError:
                        status_text.text(f"⚠️ Chunk {i+1}: Audio unclear, skipping...")
                        break
                        
                    except (sr.RequestError, TimeoutError, Exception) as e:
                        if attempt < max_retries - 1:
                            status_text.text(f"⚠️ Chunk {i+1}: Retrying... (Attempt {attempt + 2}/{max_retries})")
                            time.sleep(retry_delay)
                            retry_delay *= 2
                        else:
                            st.warning(f"❌ Chunk {i+1} failed: Network issue. Try AssemblyAI instead.")
                            failed_chunks.append(i+1)
                            break
                
                if text:
                    if show_timestamps:
                        minutes = (i * chunk_length_ms) // 60000
                        seconds = ((i * chunk_length_ms) % 60000) // 1000
                        timestamp = f"({minutes:02d}:{seconds:02d})"
                        timestamped_text.append({"time": timestamp, "text": text})
                        full_text += f"{timestamp} {text} "
                    else:
                        full_text += text + " "
            
            if os.path.exists(chunk_filename):
                os.remove(chunk_filename)
            
            time.sleep(0.5)
        
        if os.path.exists(temp_wav):
            os.remove(temp_wav)
        
        if failed_chunks:
            status_text.text(f"⚠️ Transcription complete! Failed chunks: {len(failed_chunks)}/{len(chunks)}")
            st.info("💡 For better results, try using AssemblyAI (Cloud - FASTEST)")
        else:
            status_text.text("✅ Transcription complete!")
        
        # if not full_text.strip():
        #     st.error("❌ Google API failed to transcribe. Please use AssemblyAI or Vosk instead.")
        #     return None, None
        
        if show_timestamps:
            return full_text.strip(), timestamped_text
        return full_text.strip(), None
        
    except Exception as e:
        st.error(f"Google API Error: {e}")
        st.info("💡 Try using AssemblyAI (fastest) or Vosk (offline) instead")
        return None, None

# Export functions
def export_summary_to_pdf(summary_text, filename):
    """Export formatted summary to PDF"""
    if not PDF_AVAILABLE:
        return summary_text.encode('utf-8')
    
    try:
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, 'AI Summary', 0, 1, 'C')
        pdf.ln(5)
        
        pdf.set_font("Arial", size=10)
        
        lines = summary_text.split('\n')
        for line in lines:
            safe_line = line.encode('ascii', 'ignore').decode('ascii')
            if safe_line.startswith('#'):
                pdf.set_font("Arial", 'B', 12)
                pdf.multi_cell(0, 6, txt=safe_line.replace('#', '').strip())
                pdf.set_font("Arial", size=10)
            else:
                pdf.multi_cell(0, 6, txt=safe_line)
            pdf.ln(2)
        
        return bytes(pdf.output())
        
    except Exception as e:
        st.error(f"Error creating PDF: {e}")
        return summary_text.encode('utf-8')

def export_summary_to_docx(summary_text, filename):
    """Export formatted summary to DOCX"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        title = doc.add_heading('AI Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lines = summary_text.split('\n')
        for line in lines:
            if line.startswith('#'):
                heading_text = line.replace('#', '').strip()
                doc.add_heading(heading_text, level=2)
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.getvalue()
    except Exception as e:
        st.error(f"Error creating DOCX: {e}")
        return None

def export_to_pdf(text, filename):
    """Export transcription to PDF"""
    try:
        if PDF_AVAILABLE:
            try:
                from fpdf import FPDF
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                
                pdf.set_font("Arial", 'B', 16)
                pdf.cell(0, 10, 'Transcription', 0, 1, 'C')
                pdf.ln(10)
                
                pdf.set_font("Arial", size=11)
                
                clean_text = text.replace('\r', '').replace('\x00', '')
                paragraphs = clean_text.split('\n')
                
                for para in paragraphs:
                    if para.strip():
                        safe_para = para.encode('ascii', 'ignore').decode('ascii')
                        pdf.multi_cell(0, 6, txt=safe_para)
                        pdf.ln(2)
                
                return bytes(pdf.output())
                
            except Exception as e:
                st.warning(f"FPDF failed: {e}. Using simple text PDF.")
        
        st.info("PDF library not available. Downloading as formatted text instead.")
        return text.encode('utf-8')
        
    except Exception as e:
        st.error(f"Error creating PDF: {e}")
        return None

def export_to_docx(text, filename):
    """Export transcription to DOCX"""
    if not DOCX_AVAILABLE:
        st.error("DOCX export requires python-docx. Install it with: pip install python-docx")
        return None
    
    try:
        doc = Document()
        doc.add_heading('Transcription', 0)
        doc.add_paragraph(text)
        
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.getvalue()
    except Exception as e:
        st.error(f"Error creating DOCX: {e}")
        return None

def summarize_with_gemini(transcription_text):
    """Summarize transcription using Gemini AI"""
    if not GEMINI_AVAILABLE:
        st.error("Gemini not installed. Install it with: pip install google-generativeai")
        return None
    
    try:
        api_key = st.session_state.gemini_api_key
        
        if not api_key:
            return None
        
        genai.configure(api_key=api_key)
        
        with st.spinner("🤖 Gemini AI is analyzing your transcription..."):
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            prompt = f"""Please analyze this transcription and provide a well-formatted summary using this EXACT structure:

# 📘 Summary
[Provide 2-3 concise paragraphs summarizing the main points]

# 🔑 Key Points
- [Key point 1]
- [Key point 2]
- [Key point 3]
[Add more as needed]

# 📝 Action Items
- [Action item 1 - if applicable]
- [Action item 2 - if applicable]
[Or write "No specific action items mentioned" if none]

# 💬 Important Quotes
- "[Quote 1]"
- "[Quote 2]"
- "[Quote 3]"
[Select 2-3 most notable quotes]

Transcription:
{transcription_text}

Please maintain the exact heading format with emojis."""

            response = model.generate_content(prompt)
            
            return response.text
            
    except Exception as e:
        st.error(f"Gemini API Error: {e}")
        
        if "API_KEY_INVALID" in str(e) or "invalid" in str(e).lower():
            st.warning("❌ Invalid API key. Please check and re-enter.")
            if 'gemini_api_key' in st.session_state:
                del st.session_state.gemini_api_key
        
        return None

# ==================== SOW & MOM GENERATION FUNCTIONS ====================

def generate_sow_from_transcription(transcription_text, api_key):
    """Generate Statement of Work from transcription using Gemini AI"""
    if not GEMINI_AVAILABLE:
        st.error("Gemini not installed. Install it with: pip install google-generativeai")
        return None
    
    try:
        genai.configure(api_key=api_key)
        
        with st.spinner("🤖 Generating Statement of Work..."):
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            prompt = f"""You are an expert RPA consultant. Based on the following meeting/process transcription, create a professional Statement of Work (SOW) document.

TRANSCRIPTION:
{transcription_text}

OUTPUT FORMAT (REQUIRED) — SOW:

Process 1: <Process Name>

This process automates <one- to two-line description of what is automated and the business outcome>.

- Trigger: <one sentence stating the initiating event. Keep the "Trigger:" label EXACTLY>

- Data Extraction: <what data/files are ingested, where from, and what fields are extracted. If not about files, state what data is captured from forms/APIs/voice. Keep the "Data Extraction:" label EXACTLY>

- <System Entry Label>: <describe the main system interaction, e.g., "SAP Entry", "CRM Entry", "Core Banking Entry", etc. If unclear, use "System Entry". Explain the navigation/location and the key fields/actions performed.>

- System Update: <what the system updates/creates (records, statuses), validations, approvals, saves, or notifications.>

- Outcome: <final state and how downstream business logic or policies use these values; mention accuracy/consistency benefits if stated.>

Effort Table

Activity | Discover | Design | Develop | Debug | Deploy | Drive | Documentation | Project Management | Total
<Process Short Name> | <int 1–3> | <int 1–3> | <int 3–8> | <int 1–3> | <int 1–3> | <int 1–2> | <int 1–2> | <int 1–2> | <sum>

EFFORT RULES:
- All values are whole numbers. Total must equal the sum.
- Use typical band 12–20 unless the transcript clearly indicates simpler or more complex work.
- If uncertainty is high, bias Discover/Design upward by +1 and note TBDs in prose above (not as a separate section).

STYLE RULES:
- Keep sentences short, neutral, and factual.
- Use the exact bullets and bold labels shown above (do not add new sections).
- Do not invent specifics; use "TBD" where details are missing.
- Maintain the same capitalization for headings as shown.

Extract all relevant information from the transcription and create the SOW following this EXACT format."""

            response = model.generate_content(prompt)
            return response.text
            
    except Exception as e:
        st.error(f"SOW Generation Error: {e}")
        return None

def generate_mom_from_transcription(transcription_text, api_key):
    """Generate Minutes of Meeting from transcription using Gemini AI"""
    if not GEMINI_AVAILABLE:
        st.error("Gemini not installed. Install it with: pip install google-generativeai")
        return None
    
    try:
        genai.configure(api_key=api_key)
        
        with st.spinner("🤖 Generating Minutes of Meeting..."):
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            prompt = f"""You are an expert meeting coordinator. Based on the following meeting transcription, create professional Minutes of Meeting (MoM).

TRANSCRIPTION:
{transcription_text}

OUTPUT FORMAT (REQUIRED) — Minutes of Meeting:

Meeting Details:
- Meeting Title: <TBD or inferred from transcript>
- Date: <Date of meeting or TBD>
- Time: <Time or TBD>
- Mode: <In‑person / Online / Call / TBD>
- Participants: <List of names or roles mentioned; use TBD if unclear>

Agenda:
- <Agenda item 1>
- <Agenda item 2>
- <Add only what is explicitly discussed>

Discussion Summary:
- <Key discussion point 1 written factually>
- <Key discussion point 2>
- <Avoid opinions or assumptions>

Decisions Made:
- <Decision 1>
- <Decision 2>
- Write "None" if no decisions were made.

Action Items:
Action | Owner | Due Date | Status
<Description> | <Name/Role or TBD> | <Date or TBD> | Open

Risks / Dependencies (if discussed):
- <Risk or dependency mentioned>
- Write "None" if not discussed.

Next Steps:
- <Immediate next step agreed>
- <Follow‑up meeting or deliverable if mentioned>

Next Meeting:
- Date: <Date or TBD>
- Time: <Time or TBD>
- Purpose: <Purpose or TBD>

STYLE & FORMATTING RULES:
- Use bullet points only (no paragraphs).
- Keep sentences short and precise.
- Do not add sections beyond those listed.
- Maintain professional SOW/MoM tone.
- If multiple meetings are discussed, create ONE MoM for the main meeting only.

Extract all relevant information from the transcription and create the MoM following this EXACT format."""

            response = model.generate_content(prompt)
            return response.text
            
    except Exception as e:
        st.error(f"MoM Generation Error: {e}")
        return None

def export_sow_to_docx(sow_text, filename="SOW_Document.docx"):
    """Export SOW to DOCX format"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Scope of Work (SOW)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # Parse and format content
        lines = sow_text.split('\n')
        in_table = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a table header
            if 'Activity |' in line and 'Discover' in line:
                in_table = True
                # Create table
                headers = [h.strip() for h in line.split('|')]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                continue
            
            # Check if line is table data
            if in_table and '|' in line:
                cells_data = [c.strip() for c in line.split('|')]
                row_cells = table.add_row().cells
                for i, data in enumerate(cells_data):
                    if i < len(row_cells):
                        row_cells[i].text = data
                in_table = False
                continue
            
            # Process heading
            if line.startswith('Process ') or line.startswith('Effort Table'):
                doc.add_heading(line, level=1)
            elif line.startswith('•'):
                # Extract bold label if present
                if ':' in line:
                    label, content = line.split(':', 1)
                    p = doc.add_paragraph()
                    p.add_run(label + ':').bold = True
                    p.add_run(content)
                else:
                    doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.getvalue()
        
    except Exception as e:
        st.error(f"Error creating SOW DOCX: {e}")
        return None

def export_mom_to_docx(mom_text, filename="MOM_Document.docx"):
    """Export MoM to DOCX format"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Minutes of Meeting', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # Parse and format content
        lines = mom_text.split('\n')
        in_table = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a table header
            if 'Action |' in line and 'Owner' in line:
                in_table = True
                # Create table
                headers = [h.strip() for h in line.split('|')]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                continue
            
            # Check if line is table data
            if in_table and '|' in line and not line.startswith('Action |'):
                cells_data = [c.strip() for c in line.split('|')]
                row_cells = table.add_row().cells
                for i, data in enumerate(cells_data):
                    if i < len(row_cells):
                        row_cells[i].text = data
                continue
            elif in_table and '|' not in line:
                in_table = False
            
            # Process headings
            if line.endswith(':') and not line.startswith('•'):
                doc.add_heading(line.rstrip(':'), level=2)
            elif line.startswith('•'):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.getvalue()
        
    except Exception as e:
        st.error(f"Error creating MoM DOCX: {e}")
        return None

def export_to_srt(timestamped_text):
    """Export transcription to SRT subtitle format"""
    srt_content = ""
    for i, entry in enumerate(timestamped_text, 1):
        srt_content += f"{i}\n"
        srt_content += f"{entry['time']} --> {entry['time']}\n"
        srt_content += f"{entry['text']}\n\n"
    return srt_content

# Header
st.markdown("""
    <div class="main-header">
        <div class="header-content">
            <div class="logo-section">
                <div class="logo-circle">AS</div>
                <h2 style="margin: 0;">AudioScribe</h2>
            </div>
            <div style="display: flex; gap: 2rem; align-items: center;">
             <!--   <a href="#" style="color: white; text-decoration: none;">PRICING</a> -->
                <a href="#" style="color: white; text-decoration: none;">FAQS</a>
                <a href="#" style="color: white; text-decoration: none;">BLOG</a> 
                <span>👤 user@gmail.com</span>
            </div>
        </div>
    </div>
""", unsafe_allow_html=True)

# Layout
col_sidebar, col_main = st.columns([1, 4])

# Sidebar
with col_sidebar:
    st.markdown("### 📊 Usage")
    
    # Calculate usage stats
    today = datetime.now().strftime("%Y-%m-%d")
    today_count = sum(1 for t in st.session_state.transcriptions 
                     if t.get('date', '').startswith(datetime.now().strftime("%b %d, %Y")))
    total_count = len(st.session_state.transcriptions)
    
    st.markdown(f"""
    <div class="usage-stats">
        <p style="margin: 0; font-size: 0.9em; color: #666;">Today's Transcriptions</p>
        <p style="margin: 0; font-size: 1.5em; font-weight: bold; color: #009688;">{today_count}</p>
    </div>
    <div class="usage-stats">
        <p style="margin: 0; font-size: 0.9em; color: #666;">Total Saved</p>
        <p style="margin: 0; font-size: 1.5em; font-weight: bold; color: #009688;">{total_count}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # if st.button("🔄 GO UNLIMITED", use_container_width=True):
    #     st.info("Upgrade to unlimited transcriptions!")
    
    st.markdown("---")
    st.markdown("### 📁 Shortcuts")
    
    if st.button("📄 Recent Files", use_container_width=True):
        st.session_state.show_recent_files = not st.session_state.show_recent_files
        st.session_state.show_search = False
    
# Main content area

with col_main:
    # Header with buttons
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("## 📋 Recent Files")
    with col2:
        if st.button("🔍 Search", use_container_width=True):
            st.session_state.show_search = not st.session_state.show_search
            st.session_state.show_recent_files = False
    
    # Search functionality
    if st.session_state.show_search:
        with st.container():
            st.markdown("### 🔍 Search Transcriptions")
            search_query = st.text_input("Search by filename or content...", key="search_input")
            
            if search_query:
                results = []
                for trans in st.session_state.transcriptions:
                    if (search_query.lower() in trans['filename'].lower() or 
                        search_query.lower() in trans['text'].lower()):
                        results.append(trans)
                
                if results:
                    st.markdown(f"**Found {len(results)} result(s):**")
                    for idx, result in enumerate(results):
                        if st.button(f"📄 {result['filename']} - {result['date']}", key=f"search_{idx}", use_container_width=True):
                            st.session_state.current_transcription = result
                            st.session_state.show_search = False
                            st.session_state.show_upload_modal = False
                            st.rerun()
                else:
                    st.info("No results found.")
    
    # Recent Files panel
    if st.session_state.show_recent_files:
        with st.container():
            st.markdown("### 📄 Recent Files")
            
            if st.session_state.transcriptions:
                sorted_transcriptions = sorted(
                    st.session_state.transcriptions,
                    key=lambda x: x.get('date', ''),
                    reverse=True
                )
                
                for idx, trans in enumerate(sorted_transcriptions):
                    col_a, col_b = st.columns([4, 1])
                    with col_a:
                        if st.button(
                            f"📄 {trans['filename']}",
                            key=f"recent_{idx}",
                            use_container_width=True
                        ):
                            st.session_state.current_transcription = trans
                            st.session_state.show_recent_files = False
                            st.session_state.show_upload_modal = False
                            st.rerun()
                    with col_b:
                        st.caption(trans.get('date', 'N/A')[:10])
            else:
                st.info("No transcriptions saved yet.")
    
    # Show transcribe button
    if st.button("☁️ TRANSCRIBE FILES", use_container_width=True, type="primary"):
        st.session_state.show_upload_modal = True
        st.session_state.show_search = False
        st.session_state.show_recent_files = False
    
    
    
    
    # Upload Modal
    if st.session_state.show_upload_modal or st.session_state.current_transcription is None:
        with st.container():
            st.markdown('<div class="upload-modal">', unsafe_allow_html=True)
            st.markdown("### ☁️ Transcribe Files")
            
            

    
            # File uploader
            st.markdown('<div class="drag-drop-area">', unsafe_allow_html=True)
            uploaded_file = st.file_uploader(
                "Drag and drop file here\nLimit 500MB per file • MP3, MP4, M4A, MOV, AAC, WAV, OGG, OPUS, MPEG, WMA, WMV, MPG, MPEG4",
                type=['mp3', 'mp4', 'm4a', 'mov', 'aac', 'wav', 'ogg', 'opus', 'mpeg', 'wma', 'wmv', 'mpg', 'mpeg4'],
                label_visibility="collapsed",
                accept_multiple_files=False,
                key="file_uploader",
                help="File must be 500MB or smaller."
            )
            if not uploaded_file:
                st.markdown("**MP3, MP4, M4A, MOV, AAC, WAV, OGG, OPUS, MPEG, WMA, WMV, MPG, MPEG4**")
                st.markdown("— OR —")
                st.markdown("**BROWSE FILES**")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Language selection
            col1, col2 = st.columns(2)
            with col1:
                language = st.selectbox(
                    "🌍 Audio Language",
                    ["English 🇺🇸", "Hindi 🇮🇳", "Spanish 🇪🇸", "French 🇫🇷"]
                )
            
            # Additional settings
            with st.expander("👥 Speaker Recognition & More Settings"):
                show_timestamps = st.checkbox("Show Timestamps", value=True)
                speaker_recognition = st.checkbox("Enable Speaker Recognition")
                
                # Transcription engine selection
                st.markdown("**Transcription Engine:**")
                
                engines = []
                engines.append("🚀 AssemblyAI (Cloud - FASTEST)")
                engines.append("⚡ Google API (Fast - Parallel)")
                #if VOSK_AVAILABLE:
                    #engines.append("💻 Vosk (Offline - Fast)")
                
                engine = st.radio(
                    "Choose engine",
                    engines,
                    index=0,
                    help="AssemblyAI is the fastest!"
                )
            
            # Transcribe button
            if uploaded_file and st.button("🎬 TRANSCRIBE", use_container_width=True, type="primary"):
                # Save uploaded file
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    video_path = tmp_file.name
                
                with st.spinner("🎬 Extracting audio from video..."):
                    audio_path = transcribe_video_to_text_enhanced(video_path)
                
                if audio_path:
                    # Choose transcription engine
                    if 'engine' in locals():
                        if "AssemblyAI" in engine:
                            with st.spinner("☁️ Transcribing with AssemblyAI..."):
                                transcription, timestamped = transcribe_with_assemblyai(
                                    audio_path,
                                    show_timestamps=show_timestamps
                                )
                                # 
                                # video_result = get_insights_from_video(transcription, user_query)
                        # elif "Vosk" in engine and VOSK_AVAILABLE:
                        #     with st.spinner("💻 Transcribing with Vosk..."):
                        #         transcription, timestamped = transcribe_with_vosk(
                        #             audio_path,
                        #             show_timestamps=show_timestamps
                        #         )
                        else:
                            with st.spinner("⚡ Transcribing with Google API (Parallel)..."):
                                transcription, timestamped = transcribe_long_audio_enhanced(
                                    audio_path, 
                                    chunk_length_ms=180000,  # 1 minute per chunk
                                    show_timestamps=show_timestamps
                                )
                                # video_result = get_insights_from_video(transcription, user_query)
                                
                    else:
                        with st.spinner("🎙️ Transcribing with Google Speech API..."):
                            transcription, timestamped = transcribe_long_audio_enhanced(
                                audio_path, 
                                chunk_length_ms=180000,  # 1 minute per chunk
                                show_timestamps=show_timestamps
                            )
                            # video_result = get_insights_from_video(transcription, user_query)
                    
                    if transcription:
                        # Calculate duration (simple estimation)
                        try:
                            audio = AudioSegment.from_file(audio_path)
                            duration_seconds = len(audio) / 1000
                            duration_minutes = int(duration_seconds / 60)
                            duration = f"{duration_minutes}m"
                        except:
                            duration = "N/A"
                        
                        # Save transcription
                        st.session_state.current_transcription = {
                            'filename': uploaded_file.name,
                            'text': transcription,
                            'timestamped': timestamped,
                            'date': datetime.now().strftime("%b %d, %Y, %I:%M %p"),
                            'duration': duration,
                            'summary': None
                        }
                        st.session_state.transcriptions.append(st.session_state.current_transcription)
                        st.session_state.show_upload_modal = False
                        
                        # Cleanup
                        if os.path.exists(audio_path):
                            os.remove(audio_path)
                        if os.path.exists(video_path):
                            os.remove(video_path)
                        
                        st.success("✅ Transcription saved successfully!")
                        st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            
    
    # Display transcription result
    if st.session_state.current_transcription and not st.session_state.show_upload_modal:
        st.markdown("---")
        
        
        
        # ==================== CHATBOT SECTION ====================
#     if st.session_state.current_transcription and not st.session_state.show_upload_modal:
#         st.markdown("---")
        
#         # Chatbot Header
#         st.markdown("""
#             <div style='background: rgba(255,255,255,0.95); padding: 1rem 1.5rem; 
#                     border-radius: 10px; margin-top: 1rem; border-left: 4px solid #009688;'>
#             <h3 style='color:#00796B; margin:0.5rem 0;'>💬 Ask Questions About This Transcription</h3>
#             <p style='color:#374151; font-size: 14px; margin: 0.5rem 0;'>
#                 🤖 Ask anything about key points, decisions, action items, or specific details from the meeting.
#             </p>
#         </div>
#         """, unsafe_allow_html=True)
        
#         # Initialize chat history
#         if 'chat_history' not in st.session_state:
#             st.session_state.chat_history = []
        
#         # Display chat messages
#         if st.session_state.chat_history:
#             st.markdown("#### 📝 Conversation History")
            
#             for i, chat in enumerate(st.session_state.chat_history):
#                 # User message
#                 st.markdown(f"""
#                 <div style="background:#E3F2FD; padding:12px 16px; border-radius:10px; 
#                            margin:10px 0; border-left:4px solid #2196F3;">
#                     <strong style='color:#0D47A1;'>👤 You:</strong>
#                     <div style='color:#1E293B; margin-top:6px;'>{chat['question']}</div>
#                 </div>
#                 """, unsafe_allow_html=True)
                
#                 # Bot response (HTML content from chatbot)
#                 # Bot response (HTML content from chatbot)
#                 st.markdown(f"""
# <div style="background:#1B5E20; padding:12px 16px; border-radius:10px; 
#            margin:10px 0; border-left:4px solid #4CAF50;">
#     <strong style='color:#FFFFFF;'>🤖 Assistant:</strong>
#     <div style='margin-top:6px; color:#FFFFFF;'>{chat['answer']}</div>
# </div>
# """, unsafe_allow_html=True)
        
#         # Chat input form
#         st.markdown("#### 💭 Ask Your Question")
        
#         with st.form(key="chatbot_form", clear_on_submit=True):
#             user_question = st.text_area(
#                 "Type your question here:",
#                 placeholder="Example: What were the main decisions made in this meeting? What action items were discussed?",
#                 height=80,
#                 key="chat_input"
#             )
            
#             col1, col2 = st.columns([3, 1])
            
#             with col1:
#                 submit_chat = st.form_submit_button("🚀 Ask Question", type="primary", use_container_width=True)
            
#             with col2:
#                 if st.session_state.chat_history:
#                     if st.form_submit_button("🗑️ Clear Chat", use_container_width=True):
#                         st.session_state.chat_history = []
#                         st.rerun()
        
#         # Process chat query
#         if submit_chat and user_question and user_question.strip():
#             with st.spinner("🤖 Analyzing transcription and generating answer..."):
#                 try:
#                     # Call the async chatbot function
#                     answer = user_query(
#                         user_question,
#                         st.session_state.current_transcription["text"]
#                     )
                    
#                     # Save to chat history
#                     st.session_state.chat_history.append({
#                         'question': user_question,
#                         'answer': answer,
#                         'timestamp': datetime.now().strftime("%I:%M %p")
#                     })
                    
#                     st.success("✅ Answer generated!")
#                     st.rerun()
                    
#                 except Exception as e:
#                     st.error(f"❌ Error: {str(e)}")
#                     st.info("💡 Make sure your Gemini API key is configured correctly")
        
        # ==================== CHATBOT SECTION ====================
        # ==================== CHATBOT SECTION ====================
        if st.session_state.current_transcription and not st.session_state.show_upload_modal:
            st.markdown("---")

            st.markdown("""
    <div style='background: rgba(255,255,255,0.85); padding: 8px 16px; 
                border-radius: 8px; margin-top: 1rem;'>
        <h4 style='color:#00796B; margin:4px 0;'>💬 Ask Questions About This Transcription</h4>
        <p style='color:#374151; font-size: 13px; margin: 0;'>
            🤖 Ask anything about key points, decisions, issues, or attendees.
        </p>
    </div>
""", unsafe_allow_html=True)

    
            # Initialize chat history
            if 'chat_history' not in st.session_state:
                st.session_state.chat_history = []

            # Display chat messages
            if st.session_state.chat_history:
                st.markdown("#### 📝 Conversation")
        
                for i, chat in enumerate(st.session_state.chat_history):
                    # User message container
                    st.markdown(f"""
                        <div style="
                        background:#E3F2FD;
                        padding:8px 14px;
                        border-radius:8px;
                        margin:8px 0;
                        border-left:4px solid #2196F3;
                        font-size:14px;
                ">
                    <strong style='color:#0D47A1;'>👤 You:</strong>
                    <div style='color:#1E293B; margin-top:4px;'>{chat['question']}</div>
                </div>
                """, unsafe_allow_html=True)

                    # Bot response container - header
                    st.markdown(f"""
<div style="background: linear-gradient(135deg, #009688 0%, #00796B 100%); 
           padding:12px 16px; border-radius:10px; 
           margin:10px 0; border-left:4px solid #00BCD4;
           box-shadow: 0 4px 16px rgba(0, 150, 136, 0.4);">
    <strong style='color:#FFFFFF;'>🤖 Assistant:</strong>
    <div style='margin-top:6px; color:#FFFFFF !important;'>{chat['answer']}</div>
</div>
""", unsafe_allow_html=True)


            # Chat input form
            st.markdown("#### 💭 Ask Your Question")
    
            with st.form(key="chatbot_form", clear_on_submit=True):
                user_question = st.text_area(
                    "Type your question here:",
                    placeholder="Ask anything about this meeting",
                    height=50,
                    key="chat_input"
                )

                col1, col2 = st.columns([3, 1])

                with col1:
                    submit_chat = st.form_submit_button("🚀 Ask Question", type="primary", use_container_width=True)

                with col2:
                    if st.session_state.chat_history:
                        if st.form_submit_button("🗑️ Clear Chat", use_container_width=True):
                            st.session_state.chat_history = []
                            st.rerun()

            

            # Process chat query
            if submit_chat and user_question and user_question.strip():
                with st.spinner("🤖 Thinking..."):
                    try:
                        # answer = get_insights_from_video(
                        #     user_question,
                        #     st.session_state.current_transcription["text"]
                        # )
                        answer = ask_backend(
                        user_question,
                        st.session_state.current_transcription["text"]
                        )


                        st.session_state.chat_history.append({
                            'question': user_question,
                            'answer': answer,
                            'timestamp': datetime.now().strftime("%I:%M %p")
                        })

                        st.success("✅ Answer ready!")
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
    
        




        # File header
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"## 📄 {st.session_state.current_transcription['filename']}")
            st.caption(f"📅 {st.session_state.current_transcription['date']} | ⏱️ {st.session_state.current_transcription['duration']}")
        
        with col2:
            st.markdown("### 📤 Export")
            
            # Export options
            txt_data = st.session_state.current_transcription['text']
            
            # TXT export
            st.download_button(
                "📋 Download TXT",
                data=txt_data,
                file_name="transcription.txt",
                mime="text/plain",
                use_container_width=True,
                key="download_txt"
            )
            
            # PDF export
            try:
                pdf_data = export_to_pdf(txt_data, "transcription.pdf")
                if pdf_data:
                    st.download_button(
                        "📄 Download PDF",
                        data=pdf_data,
                        file_name="transcription.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="download_pdf"
                    )
            except Exception as e:
                st.button("📄 PDF Error (Use TXT)", disabled=True, use_container_width=True)
            
            # DOCX export
            if DOCX_AVAILABLE:
                try:
                    docx_data = export_to_docx(txt_data, "transcription.docx")
                    if docx_data:
                        st.download_button(
                            "📝 Download DOCX",
                            data=docx_data,
                            file_name="transcription.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="download_docx"
                        )
                except Exception as e:
                    st.button("📝 DOCX Error", disabled=True, use_container_width=True)
            else:
                st.button("📝 Download DOCX (Install python-docx)", disabled=True, use_container_width=True)
            
            # SRT export
            if st.session_state.current_transcription['timestamped']:
                srt_data = export_to_srt(st.session_state.current_transcription['timestamped'])
                st.download_button(
                    "🎬 Download SRT",
                    data=srt_data,
                    file_name="transcription.srt",
                    mime="text/plain",
                    use_container_width=True,
                    key="download_srt"
                )
            
            st.markdown("---")
            st.markdown("### ⚙️ More")
            
            show_ts = st.checkbox("✅ Show Timestamps", value=True)
            
            # SOW Generation Button
            if st.button("📋 Generate SOW\nScope of Work", use_container_width=True, key="sow_button"):
                st.session_state.show_sow_generator = True
                st.session_state.show_mom_generator = False
                st.rerun()

            # MoM Generation Button  
            if st.button("📝 Generate MoM\nMinutes of Meeting", use_container_width=True, key="mom_button"):
                st.session_state.show_mom_generator = True
                st.session_state.show_sow_generator = False
                st.rerun()
            
            # Gemini Summarization
            if GEMINI_AVAILABLE:
                # Check if API key is set
                if not st.session_state.gemini_api_key:
                    st.markdown("**🔑 Gemini AI Summary**")
                    st.info("Enter your Gemini API key to enable AI summarization")
                    
                    with st.form("gemini_key_form"):
                        api_key_input = st.text_input(
                            "Gemini API Key:", 
                            type="password",
                            help="Get FREE key at: https://makersuite.google.com/app/apikey"
                        )
                        submit = st.form_submit_button("Save & Continue", use_container_width=True)
                        
                        if submit and api_key_input:
                            st.session_state.gemini_api_key = api_key_input
                            st.success("✅ API key saved!")
                            st.rerun()
                else:
                    if st.button("🤖 Gemini AI\nSummarize with AI", use_container_width=True):
                        summary = summarize_with_gemini(st.session_state.current_transcription['text'])
                        if summary:
                            st.session_state.current_transcription['summary'] = summary
                            # Update in main list
                            for trans in st.session_state.transcriptions:
                                if trans['filename'] == st.session_state.current_transcription['filename'] and trans['date'] == st.session_state.current_transcription['date']:
                                    trans['summary'] = summary
                            st.rerun()
            else:
                st.button("🤖 Gemini AI\n(Install: pip install google-generativeai)", disabled=True, use_container_width=True)
            
            if st.button("🌐 Translate\n to 134+ languages", use_container_width=True):
                st.info("Translation feature coming soon!")
                
            # ==================== SOW GENERATION SECTION ====================
    # if st.session_state.show_sow_generator and st.session_state.current_transcription:
    #     st.markdown("---")
    #     st.markdown("---")

    #     # Header
    #     col_h1, col_h2 = st.columns([5, 1])
    #     with col_h1:
    #         st.markdown("## 📋 Scope of Work (SOW) Generator")
    #         st.info("🎯 Generate a professional SOW document from your transcription")
    #     with col_h2:
    #         if st.button("❌ Close", use_container_width=True, key="close_sow"):
    #             st.session_state.show_sow_generator = False
    #             st.rerun()

    #     st.markdown("---")

    #     # Transcription Preview
#     st.subheader("📄 Source Transcription")
#     st.info(f"✅ Using transcription from: **{st.session_state.current_transcription['filename']}**")
    
#     trans_preview = st.session_state.current_transcription['text'][:1000]
#     st.text_area(
#         "Transcription Preview (first 1000 chars)",
#         trans_preview + ("..." if len(st.session_state.current_transcription['text']) > 1000 else ""),
#         height=150,
#         disabled=True,
#         key="sow_trans_preview"
#     )

    #     # Additional context
    #     additional_sow_context = st.text_area(
    #         "Additional Context (Optional)",
    #         placeholder="Add any specific SOW requirements, process details, or constraints...",
    #         height=100,
    #         key="sow_additional_context"
    #     )

    #     st.markdown("---")

    #     # Generate SOW Button
    #     col_gen1, col_gen2 = st.columns([3, 1])

    #     with col_gen1:
    #         generate_sow_button = st.button(
    #             "🚀 Generate SOW Document",
    #             type="primary",
    #             use_container_width=True,
    #             key="generate_sow_btn"
    #         )

    #     with col_gen2:
    #         st.metric("Est. Time", "30-60s")

    #     # SOW GENERATION LOGIC
    #     if generate_sow_button:
    #         # Check API Key
    #         if not st.session_state.gemini_api_key:
    #             st.error("⚠️ Please configure your Gemini API key first")
    #             with st.form("sow_api_key_form"):
    #                 quick_key = st.text_input("Enter Gemini API Key:", type="password")
    #                 if st.form_submit_button("Save Key"):
    #                     st.session_state.gemini_api_key = quick_key
    #                     st.success("✅ API key saved! Click Generate SOW again.")
    #                     st.rerun()
    #         else:
    #             # Prepare content
    #             full_content = st.session_state.current_transcription['text']
    #             if additional_sow_context:
    #                 full_content = f"{full_content}\n\nADDITIONAL CONTEXT:\n{additional_sow_context}"

    #             # Generate SOW
    #             sow_content = generate_sow_from_transcription(full_content, st.session_state.gemini_api_key)

    #             if sow_content:
    #                 st.session_state.generated_sow = sow_content

    #                 # SUCCESS!
    #                 st.balloons()
    #                 st.success("🎉 SOW Generated Successfully!")

    #                 # Display SOW
    #                 st.markdown("### 📄 Generated Statement of Work")
    #                 st.markdown(sow_content)

    #                 st.markdown("---")
    #                 st.markdown("#### 📥 Download SOW")

    #                 col_dl1, col_dl2, col_dl3 = st.columns(3)

    #                 with col_dl1:
    #                     st.download_button(
    #                         "📋 Download TXT",
    #                         data=sow_content,
    #                         file_name=f"SOW_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
    #                         mime="text/plain",
    #                         use_container_width=True,
    #                         key="download_sow_txt"
    #                     )

    #                 with col_dl2:
    #                     if DOCX_AVAILABLE:
    #                         docx_data = export_sow_to_docx(sow_content)
    #                         if docx_data:
    #                             st.download_button(
    #                                 "📝 Download DOCX",
    #                                 data=docx_data,
    #                                 file_name=f"SOW_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
    #                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    #                                 use_container_width=True,
    #                                 key="download_sow_docx"
    #                             )

    #                 with col_dl3:
    #                     if PDF_AVAILABLE:
    #                         pdf_data = export_to_pdf(sow_content, "SOW.pdf")
    #                         if pdf_data:
    #                             st.download_button(
    #                                 "📄 Download PDF",
    #                                 data=pdf_data,
    #                                 file_name=f"SOW_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
    #                                 mime="application/pdf",
    #                                 use_container_width=True,
    #                                 key="download_sow_pdf"
    #                             )
    #             else:
    #                 st.error("❌ Failed to generate SOW. Please try again.")

    # # ==================== MOM GENERATION SECTION ====================
    # if st.session_state.show_mom_generator and st.session_state.current_transcription:
    #     st.markdown("---")
    #     st.markdown("---")

    #     # Header
    #     col_h1, col_h2 = st.columns([5, 1])
    #     with col_h1:
    #         st.markdown("## 📝 Minutes of Meeting (MoM) Generator")
    #         st.info("🎯 Generate professional meeting minutes from your transcription")
    #     with col_h2:
    #         if st.button("❌ Close", use_container_width=True, key="close_mom"):
    #             st.session_state.show_mom_generator = False
    #             st.rerun()

    #     st.markdown("---")

    #     # Transcription Preview
    #     st.subheader("📄 Source Transcription")
    #     st.info(f"✅ Using transcription from: **{st.session_state.current_transcription['filename']}**")

    #     trans_preview = st.session_state.current_transcription['text'][:1000]
    #     st.text_area(
    #         "Transcription Preview (first 1000 chars)",
    #         trans_preview + ("..." if len(st.session_state.current_transcription['text']) > 1000 else ""),
    #         height=150,
    #         disabled=True,
    #         key="mom_trans_preview"
    #     )

    #     # Meeting Details (Optional)
    #     with st.expander("📅 Meeting Details (Optional)", expanded=False):
    #         col1, col2 = st.columns(2)

    #         with col1:
    #             meeting_title = st.text_input("Meeting Title", key="mom_meeting_title")
    #             meeting_date = st.date_input("Meeting Date", key="mom_meeting_date")

    #         with col2:
    #             meeting_time = st.time_input("Meeting Time", key="mom_meeting_time")
    #             meeting_mode = st.selectbox("Meeting Mode", ["Online", "In-person", "Call", "Hybrid"], key="mom_meeting_mode")

    #     # Additional context
    #     additional_mom_context = st.text_area(
    #         "Additional Context (Optional)",
    #         placeholder="Add participant names, specific action items, or other details...",
    #         height=100,
    #         key="mom_additional_context"
    #     )

    #     st.markdown("---")

    #     # Generate MoM Button
    #     col_gen1, col_gen2 = st.columns([3, 1])

    #     with col_gen1:
    #         generate_mom_button = st.button(
    #             "🚀 Generate MoM Document",
    #             type="primary",
    #             use_container_width=True,
    #             key="generate_mom_btn"
    #         )

    #     with col_gen2:
    #         st.metric("Est. Time", "30-60s")

    #     # MOM GENERATION LOGIC
    #     if generate_mom_button:
    #         # Check API Key
    #         if not st.session_state.gemini_api_key:
    #             st.error("⚠️ Please configure your Gemini API key first")
    #             with st.form("mom_api_key_form"):
    #                 quick_key = st.text_input("Enter Gemini API Key:", type="password")
    #                 if st.form_submit_button("Save Key"):
    #                     st.session_state.gemini_api_key = quick_key
    #                     st.success("✅ API key saved! Click Generate MoM again.")
    #                     st.rerun()
    #         else:
    #             # Prepare content
    #             full_content = st.session_state.current_transcription['text']

    #             # Add meeting details if provided
    #             if 'meeting_title' in locals() and meeting_title:
    #                 full_content = f"Meeting Title: {meeting_title}\n\n{full_content}"
    #             if 'meeting_date' in locals() and meeting_date:
    #                 full_content = f"Date: {meeting_date}\n{full_content}"
    #             if 'meeting_time' in locals() and meeting_time:
    #                 full_content = f"Time: {meeting_time}\n{full_content}"
    #             if 'meeting_mode' in locals() and meeting_mode:
    #                 full_content = f"Mode: {meeting_mode}\n{full_content}"

    #             if additional_mom_context:
    #                 full_content = f"{full_content}\n\nADDITIONAL CONTEXT:\n{additional_mom_context}"

    #             # Generate MoM
    #             mom_content = generate_mom_from_transcription(full_content, st.session_state.gemini_api_key)

    #             if mom_content:
    #                 st.session_state.generated_mom = mom_content

    #                 # SUCCESS!
    #                 st.balloons()
    #                 st.success("🎉 Minutes of Meeting Generated Successfully!")

    #                 # Display MoM
    #                 st.markdown("### 📄 Generated Minutes of Meeting")
    #                 st.markdown(mom_content)

    #                 st.markdown("---")
    #                 st.markdown("#### 📥 Download MoM")

    #                 col_dl1, col_dl2, col_dl3 = st.columns(3)

    #                 with col_dl1:
    #                     st.download_button(
    #                         "📋 Download TXT",
    #                         data=mom_content,
    #                         file_name=f"MoM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
    #                         mime="text/plain",
    #                         use_container_width=True,
    #                         key="download_mom_txt"
    #                     )

    #                 with col_dl2:
    #                     if DOCX_AVAILABLE:
    #                         docx_data = export_mom_to_docx(mom_content)
    #                         if docx_data:
    #                             st.download_button(
    #                                 "📝 Download DOCX",
    #                                 data=docx_data,
    #                                 file_name=f"MoM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
    #                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    #                                 use_container_width=True,
    #                                 key="download_mom_docx"
    #                             )

    #                 with col_dl3:
    #                     if PDF_AVAILABLE:
    #                         pdf_data = export_to_pdf(mom_content, "MoM.pdf")
    #                         if pdf_data:
    #                             st.download_button(
    #                                 "📄 Download PDF",
    #                                 data=pdf_data,
    #                                 file_name=f"MoM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
    #                                 mime="application/pdf",
    #                                 use_container_width=True,
    #                                 key="download_mom_pdf"
    #                             )
    #             else:
    #                 st.error("❌ Failed to generate MoM. Please try again.")
        
        # Transcription display
        st.markdown('<div class="transcript-container">', unsafe_allow_html=True)
        
        # Show Gemini summary if available
        if st.session_state.current_transcription.get('summary'):
            st.markdown("### 🤖 AI Summary")
            st.markdown(st.session_state.current_transcription['summary'])
            
            # Download summary buttons
            st.markdown("#### 📥 Download Summary")
            
            col_a, col_b, col_c = st.columns(3)
            
            with col_a:
                st.download_button(
                    "📋 TXT",
                    data=st.session_state.current_transcription['summary'],
                    file_name="summary.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="summary_txt"
                )
            
            with col_b:
                pdf_summary = export_summary_to_pdf(st.session_state.current_transcription['summary'], "summary.pdf")
                if pdf_summary:
                    st.download_button(
                        "📄 PDF",
                        data=pdf_summary,
                        file_name="summary.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="summary_pdf"
                    )
            
            with col_c:
                if DOCX_AVAILABLE:
                    docx_summary = export_summary_to_docx(st.session_state.current_transcription['summary'], "summary.docx")
                    if docx_summary:
                        st.download_button(
                            "📝 DOCX",
                            data=docx_summary,
                            file_name="summary.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="summary_docx"
                        )
            
            st.markdown("---")
            st.markdown("### 📝 Full Transcription")
        
        if show_ts and st.session_state.current_transcription['timestamped']:
            for entry in st.session_state.current_transcription['timestamped']:
                st.markdown(f"<span class='timestamp'>{entry['time']}</span> {entry['text']}", unsafe_allow_html=True)
                st.markdown("")
        else:
            st.markdown(st.session_state.current_transcription['text'])
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Audio player (only show if we have the uploaded file)
        # Note: Audio playback is only available during initial upload session
        st.markdown("---")
        st.info("💡 Audio playback is available only during the upload session")

# Footer
st.markdown("---")
st.markdown("""
    <div style='text-align: center; color: #6B7280; padding: 2rem;'>
        <p>©2026 AudioScribe</p>
        <p>
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>Home</a> |
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>Blog</a> |
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>Pricing</a> |
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>FAQs</a> |
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>Support</a> |
            <a href='#' style='color: #6B7280; margin: 0 0.5rem;'>Privacy</a>
        </p>
    </div>
""", unsafe_allow_html=True)
